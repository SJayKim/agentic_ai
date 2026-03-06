"""
Utility Tools — 파일/시스템 정보 조회 도구

Knowledge Graph에 의존하지 않는 가벼운 유틸리티 도구 모음.
폴더 파일 목록, 문서 메타정보, 시스템 상태 등을 직접 조회합니다.
"""

import os
import json
import re
from pathlib import Path
from typing import Dict, Any, List
from datetime import datetime

import requests as _requests
from bs4 import BeautifulSoup as _BS

from src.rag.summary_store import get_summary, get_all_summaries_brief
from src.rag.document_converter import convert_document


# 문서 저장 경로
_DOCS_DIR = Path(__file__).parent.parent.parent / "data" / "documents"
_RAG_STORAGE_DIR = Path(__file__).parent.parent.parent / "data" / "rag_storage"


def _format_file_size(size_bytes: int) -> str:
    """바이트를 읽기 쉬운 형식으로 변환"""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"


def _get_file_type(ext: str) -> str:
    """확장자를 사람이 읽기 쉬운 타입으로 변환"""
    type_map = {
        ".pdf": "PDF 문서",
        ".docx": "Word 문서",
        ".doc": "Word 문서 (구형)",
        ".pptx": "PowerPoint 프레젠테이션",
        ".ppt": "PowerPoint (구형)",
        ".xlsx": "Excel 스프레드시트",
        ".xls": "Excel (구형)",
        ".txt": "텍스트 파일",
        ".md": "마크다운 파일",
        ".csv": "CSV 데이터",
        ".json": "JSON 데이터",
    }
    return type_map.get(ext.lower(), f"{ext} 파일")


# ============ Tool Functions ============

async def list_documents() -> str:
    """
    문서 폴더에 있는 모든 파일의 목록을 반환합니다.
    파일명, 크기, 타입 정보를 포함합니다.
    """
    if not _DOCS_DIR.exists():
        return "문서 폴더가 존재하지 않습니다."

    files = []
    for f in sorted(_DOCS_DIR.iterdir()):
        if f.is_file() and not f.name.startswith("."):
            stat = f.stat()
            files.append({
                "filename": f.name,
                "type": _get_file_type(f.suffix),
                "size": _format_file_size(stat.st_size),
                "size_bytes": stat.st_size,
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
            })

    if not files:
        return "문서 폴더에 파일이 없습니다."

    # 읽기 쉬운 형식으로 출력
    lines = [f"📁 문서 폴더: {len(files)}개 파일\n"]
    total_size = 0
    for f in files:
        lines.append(f"  • {f['filename']}")
        lines.append(f"    타입: {f['type']} | 크기: {f['size']} | 수정일: {f['modified']}")
        total_size += f["size_bytes"]
    lines.append(f"\n총 용량: {_format_file_size(total_size)}")

    return "\n".join(lines)


async def get_document_info(filename: str) -> str:
    """
    특정 문서의 상세 메타정보를 반환합니다.
    파일명이 정확히 일치하지 않으면 부분 매칭을 시도합니다.
    """
    if not _DOCS_DIR.exists():
        return "문서 폴더가 존재하지 않습니다."

    # 정확한 파일명 매칭
    target = _DOCS_DIR / filename
    if not target.exists():
        # 부분 매칭 시도 (대소문자 무시)
        filename_lower = filename.lower()
        matches = []
        for f in _DOCS_DIR.iterdir():
            if f.is_file() and filename_lower in f.name.lower():
                matches.append(f)

        if not matches:
            available = [f.name for f in _DOCS_DIR.iterdir() if f.is_file()]
            return f"'{filename}' 파일을 찾을 수 없습니다.\n사용 가능한 파일: {', '.join(available)}"
        elif len(matches) == 1:
            target = matches[0]
        else:
            return f"여러 파일이 매칭됩니다: {', '.join(m.name for m in matches)}\n정확한 파일명을 지정해주세요."

    stat = target.stat()
    ext = target.suffix.lower()

    info = {
        "파일명": target.name,
        "타입": _get_file_type(ext),
        "크기": _format_file_size(stat.st_size),
        "수정일": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
        "경로": str(target),
    }

    # 파일 타입별 추가 정보
    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(str(target))
            page_count = len(doc)
            doc.close()
            info["페이지 수"] = f"{page_count}페이지"
        elif ext == ".docx":
            from docx import Document
            doc = Document(str(target))
            info["단락 수"] = f"{len(doc.paragraphs)}개"
            info["테이블 수"] = f"{len(doc.tables)}개"
        elif ext == ".pptx":
            from pptx import Presentation
            prs = Presentation(str(target))
            info["슬라이드 수"] = f"{len(prs.slides)}장"
        elif ext in (".txt", ".md"):
            with open(target, "r", encoding="utf-8", errors="replace") as f:
                content = f.read()
            info["글자 수"] = f"{len(content)}자"
            info["줄 수"] = f"{content.count(chr(10)) + 1}줄"
    except Exception as e:
        info["추가정보"] = f"분석 실패: {str(e)}"

    lines = [f"📄 문서 정보: {info['파일명']}\n"]
    for key, value in info.items():
        if key != "파일명":
            lines.append(f"  {key}: {value}")

    return "\n".join(lines)


async def get_system_status() -> str:
    """
    Knowledge Graph 시스템의 현재 상태를 반환합니다.
    노드/엣지 수, 인덱싱된 문서 수, 저장소 크기 등을 포함합니다.
    """
    status = {}

    # KG 노드/엣지 수
    graphml_path = _RAG_STORAGE_DIR / "graph_chunk_entity_relation.graphml"
    if graphml_path.exists():
        with open(graphml_path, "r", encoding="utf-8") as f:
            content = f.read()
        status["노드 수"] = len(re.findall(r'<node ', content))
        status["엣지 수"] = len(re.findall(r'<edge ', content))
    else:
        status["KG 상태"] = "그래프 파일 없음"

    # 인덱싱된 문서 수
    docs_json = _RAG_STORAGE_DIR / "kv_store_full_docs.json"
    if docs_json.exists():
        with open(docs_json, "r", encoding="utf-8") as f:
            docs = json.load(f)
        status["인덱싱된 문서 수"] = f"{len(docs)}개"

        # 문서 이름 추출
        doc_names = []
        for k, v in docs.items():
            content = v.get("content", "") if isinstance(v, dict) else str(v)
            m = re.search(r'\[Document: (.+?)\]', content)
            doc_names.append(m.group(1) if m else k[:40])
        status["인덱싱된 문서 목록"] = doc_names

    # 원본 문서 폴더 상태
    if _DOCS_DIR.exists():
        file_count = len([f for f in _DOCS_DIR.iterdir() if f.is_file() and not f.name.startswith(".")])
        status["원본 문서 수"] = f"{file_count}개"

    # 저장소 크기
    if _RAG_STORAGE_DIR.exists():
        total_size = sum(
            f.stat().st_size for f in _RAG_STORAGE_DIR.rglob("*") if f.is_file()
        )
        status["RAG 저장소 크기"] = _format_file_size(total_size)

    lines = ["📊 시스템 상태 리포트\n"]
    for key, value in status.items():
        if isinstance(value, list):
            lines.append(f"  {key}:")
            for item in value:
                lines.append(f"    • {item}")
        else:
            lines.append(f"  {key}: {value}")

    return "\n".join(lines)


async def get_document_summary(filename: str) -> str:
    """
    특정 문서의 구조화 요약을 반환합니다.
    인제스트 시 생성된 LLM 요약을 캐시에서 즉시 조회합니다.
    """
    result = get_summary(filename)

    if result is None:
        # 사용 가능한 요약 목록 제공
        all_briefs = get_all_summaries_brief()
        if all_briefs:
            available = ", ".join(b["filename"] for b in all_briefs)
            return f"'{filename}'에 대한 요약이 없습니다.\n요약이 있는 문서: {available}"
        return f"'{filename}'에 대한 요약이 없습니다. 아직 인덱싱된 문서가 없습니다."

    if isinstance(result, dict) and result.get("error") == "multiple_matches":
        matches = ", ".join(result["matches"])
        return f"여러 문서가 매칭됩니다: {matches}\n정확한 파일명을 지정해주세요."

    # 정상 결과
    lines = [
        f"📋 문서 요약: {result['filename']}",
        f"   문서 유형: {result.get('doc_type', 'N/A')}",
        f"   핵심 주제: {result.get('key_topic', 'N/A')}",
        f"   원본 크기: {result['original_chars']:,}자 → 요약: {result['summary_chars']:,}자",
        f"   인덱싱 시각: {result.get('indexed_at', 'N/A')}",
        f"\n{'='*60}\n",
        result["summary"],
    ]
    return "\n".join(lines)


async def list_document_summaries() -> str:
    """
    인덱싱된 모든 문서의 간략 요약 목록을 반환합니다.
    각 문서의 파일명, 문서 유형, 핵심 주제를 포함합니다.
    """
    briefs = get_all_summaries_brief()

    if not briefs:
        return "저장된 문서 요약이 없습니다. 문서를 인제스트하면 자동으로 요약이 생성됩니다."

    lines = [f"📚 문서 요약 목록: {len(briefs)}개 문서\n"]
    for i, b in enumerate(briefs, 1):
        lines.append(f"  {i}. {b['filename']}")
        if b.get("doc_type"):
            lines.append(f"     유형: {b['doc_type']}")
        if b.get("key_topic"):
            lines.append(f"     주제: {b['key_topic']}")
        if b.get("indexed_at"):
            lines.append(f"     인덱싱: {b['indexed_at'][:19]}")
        lines.append("")

    lines.append("상세 요약을 보려면 get_document_summary 도구를 사용하세요.")
    return "\n".join(lines)


async def get_document_content(filename: str) -> str:
    """
    특정 문서의 원문 텍스트를 추출하여 반환합니다.
    document_converter를 사용하여 실시간으로 텍스트를 추출합니다.
    """
    if not _DOCS_DIR.exists():
        return "문서 폴더가 존재하지 않습니다."

    # 정확한 파일명 매칭
    target = _DOCS_DIR / filename
    if not target.exists():
        # 부분 매칭 시도 (대소문자 무시)
        filename_lower = filename.lower()
        matches = [f for f in _DOCS_DIR.iterdir()
                   if f.is_file() and filename_lower in f.name.lower()]

        if not matches:
            available = [f.name for f in _DOCS_DIR.iterdir() if f.is_file()]
            return f"'{filename}' 파일을 찾을 수 없습니다.\n사용 가능한 파일: {', '.join(available)}"
        elif len(matches) == 1:
            target = matches[0]
        else:
            return f"여러 파일이 매칭됩니다: {', '.join(m.name for m in matches)}\n정확한 파일명을 지정해주세요."

    try:
        doc = convert_document(str(target))
        if doc.error:
            return f"문서 변환 실패: {doc.error}"
        if not doc.content.strip():
            return f"'{doc.filename}'에서 텍스트를 추출할 수 없습니다."

        # 메타데이터 헤더 + 원문
        header = (
            f"📄 원문 텍스트: {doc.filename}\n"
            f"   타입: {doc.file_type} | 페이지: {doc.page_count} | 글자 수: {len(doc.content):,}자\n"
            f"\n{'='*60}\n"
        )

        # 원문이 너무 길면 앞부분만 (LLM context 절약)
        MAX_CHARS = 15000
        content = doc.content
        if len(content) > MAX_CHARS:
            content = content[:MAX_CHARS] + f"\n\n... (총 {len(doc.content):,}자 중 앞 {MAX_CHARS:,}자만 표시)"

        return header + content
    except Exception as e:
        return f"원문 추출 중 오류: {str(e)}"


async def web_search(query: str, max_results: int = 5) -> str:
    """
    웹 검색을 수행합니다 (DuckDuckGo HTML → Naver 폴백).
    실시간 뉴스, 최신 정보, 외부 지식이 필요할 때 사용합니다.
    """
    max_results = min(max_results, 10)

    # ── 1차: DuckDuckGo HTML POST (한국어 결과 최적) ──
    results = _search_ddg_html(query, max_results)

    # ── 2차 폴백: Naver 검색 ──
    if not results:
        results = _search_naver(query, max_results)

    if not results:
        return f"'{query}'에 대한 검색 결과가 없습니다. 검색어를 변경해 다시 시도하세요."

    lines = [f"🔍 웹 검색 결과: '{query}' ({len(results)}건)\n"]
    for i, r in enumerate(results, 1):
        title = r.get("title", "제목 없음")
        body = r.get("body", "")
        href = r.get("href", "")
        lines.append(f"  {i}. **{title}**")
        if body:
            lines.append(f"     {body[:400]}")
        if href:
            lines.append(f"     출처: {href}")
        lines.append("")

    return "\n".join(lines)


# ── 검색 엔진 구현부 ──

_SEARCH_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    ),
    "Accept-Language": "ko-KR,ko;q=0.9,en-US;q=0.8,en;q=0.7",
}
_SEARCH_TIMEOUT = 12


def _search_ddg_html(query: str, max_results: int) -> list:
    """DuckDuckGo HTML POST 방식 검색 (한국어 결과 양호)."""
    try:
        resp = _requests.post(
            "https://html.duckduckgo.com/html/",
            data={"q": query, "kl": "kr-kr"},
            headers={**_SEARCH_HEADERS, "Referer": "https://html.duckduckgo.com/"},
            timeout=_SEARCH_TIMEOUT,
        )
        if resp.status_code != 200:
            return []

        soup = _BS(resp.text, "html.parser")
        results = []
        for item in soup.select(".result"):
            title_el = item.select_one(".result__a")
            snippet_el = item.select_one(".result__snippet")
            url_el = item.select_one(".result__url")
            if not title_el:
                continue
            href = ""
            if url_el:
                href = url_el.get("href", "").strip() or url_el.get_text().strip()
            if not href and title_el.get("href"):
                href = title_el["href"]
            # DDG wraps links: //duckduckgo.com/l/?uddg=<real_url>
            if "uddg=" in href:
                import urllib.parse
                parsed = urllib.parse.parse_qs(urllib.parse.urlparse(href).query)
                href = parsed.get("uddg", [href])[0]
            results.append({
                "title": title_el.get_text().strip(),
                "body": snippet_el.get_text().strip() if snippet_el else "",
                "href": href,
            })
            if len(results) >= max_results:
                break
        return results
    except Exception as e:
        print(f"[WEB-SEARCH] DDG HTML error: {e}")
        return []


def _search_naver(query: str, max_results: int) -> list:
    """Naver 웹 검색 폴백."""
    try:
        resp = _requests.get(
            "https://search.naver.com/search.naver",
            params={"where": "web", "query": query, "ie": "utf8"},
            headers=_SEARCH_HEADERS,
            timeout=_SEARCH_TIMEOUT,
        )
        if resp.status_code != 200:
            return []

        soup = _BS(resp.text, "html.parser")
        results = []
        # Naver uses dynamic class names; find links with known patterns
        for el in soup.find_all("a", class_=True):
            text = el.get_text().strip()
            href = el.get("href", "")
            cls = " ".join(el.get("class", []))
            # Skip navigation, UI elements
            if not text or len(text) < 8 or not href.startswith("http"):
                continue
            # Skip Naver internal links
            if "search.naver.com" in href or "help.naver.com" in href:
                continue
            # Look for title-like links (>8 chars, external)
            if any(kw in text for kw in query.split()[:2]) or "t_4gra5eqmK" in cls:
                # Get next sibling description
                parent = el.find_parent()
                desc = ""
                if parent:
                    desc_el = parent.find_next_sibling()
                    if desc_el:
                        desc = desc_el.get_text().strip()[:400]
                # Avoid duplicates
                if not any(r["href"] == href for r in results):
                    results.append({
                        "title": text[:200],
                        "body": desc,
                        "href": href,
                    })
            if len(results) >= max_results:
                break
        return results
    except Exception as e:
        print(f"[WEB-SEARCH] Naver error: {e}")
        return []


# ============ Tool Registration ============

def get_utility_tools_descriptions() -> List[Dict[str, Any]]:
    """유틸리티 도구의 설명 목록 반환 (Actor가 도구를 선택할 때 참조)"""
    return [
        {
            "name": "list_documents",
            "description": (
                "문서 폴더에 저장된 파일 목록을 조회합니다. "
                "각 파일의 이름, 타입(PDF/Word/PPT 등), 크기, 수정일을 반환합니다. "
                "사용 시점: '어떤 파일이 있어?', '문서 목록 보여줘', '폴더에 뭐 있어?' 등 "
                "파일 목록이나 개수를 물어볼 때 사용합니다. "
                "Knowledge Graph 검색 없이 실제 파일 시스템에서 직접 조회합니다."
            ),
            "args": {},
        },
        {
            "name": "get_document_info",
            "description": (
                "특정 문서의 상세 메타정보를 조회합니다. "
                "파일 크기, 수정일, 페이지 수(PDF), 슬라이드 수(PPT), 단락 수(Word) 등을 반환합니다. "
                "부분 파일명으로도 검색 가능합니다 (예: '영수증' → 매칭되는 파일 검색). "
                "사용 시점: '이 파일 정보 알려줘', '파일 크기가 얼마야?', '몇 페이지야?' 등 "
                "특정 문서의 물리적 속성을 물어볼 때 사용합니다. "
                "문서의 내용을 검색하려면 query_knowledge_graph를 사용하세요."
            ),
            "args": {
                "filename": {
                    "type": "string",
                    "description": "조회할 파일명 (전체 또는 부분 매칭 가능)",
                },
            },
        },
        {
            "name": "get_document_summary",
            "description": (
                "특정 문서의 구조화 요약을 조회합니다. "
                "인제스트 시 LLM이 생성한 구조화 요약(문서 유형, 핵심 주제, 주요 내용, 핵심 엔티티/관계)을 반환합니다. "
                "부분 파일명으로도 검색 가능합니다. "
                "사용 시점: '이 문서 요약해줘', '이 문서가 뭐야?', '이 파일 내용이 뭐야?', "
                "'문서 핵심 내용 알려줘' 등 문서의 전체적인 내용 파악이 필요할 때 사용합니다. "
                "query_knowledge_graph보다 빠르고 문서 전체를 포괄하는 개요를 제공합니다. "
                "특정 세부 정보나 엔티티 간 관계를 검색할 때는 query_knowledge_graph를 사용하세요."
            ),
            "args": {
                "filename": {
                    "type": "string",
                    "description": "요약을 조회할 파일명 (전체 또는 부분 매칭 가능)",
                },
            },
        },
        {
            "name": "get_document_content",
            "description": (
                "특정 문서의 원문 텍스트를 추출하여 반환합니다. "
                "PDF, Word, PPT, Excel, TXT 등 원본 파일에서 실시간으로 텍스트를 추출합니다. "
                "부분 파일명으로도 검색 가능합니다. "
                "사용 시점: '원문 보여줘', '원본 텍스트 가져와', '이 문서 전체 내용 보여줘', "
                "'문서 raw 텍스트', '원문 읽어줘' 등 문서의 가공되지 않은 원본 텍스트가 필요할 때 사용합니다. "
                "구조화 요약(get_document_summary)이나 KG 검색(query_knowledge_graph)과 달리 "
                "원본 그대로의 텍스트를 반환합니다."
            ),
            "args": {
                "filename": {
                    "type": "string",
                    "description": "원문을 조회할 파일명 (전체 또는 부분 매칭 가능)",
                },
            },
        },
        {
            "name": "list_document_summaries",
            "description": (
                "인덱싱된 모든 문서의 간략 요약 목록을 조회합니다. "
                "각 문서의 파일명, 문서 유형, 핵심 주제를 반환합니다. "
                "사용 시점: '인덱싱된 문서들 요약 보여줘', '각 문서가 뭐에 대한 거야?', "
                "'문서별 주제 알려줘' 등 전체 문서의 개요를 파악할 때 사용합니다. "
                "개별 문서의 상세 요약은 get_document_summary를 사용하세요."
            ),
            "args": {},
        },
        {
            "name": "get_system_status",
            "description": (
                "Knowledge Graph 시스템의 현재 상태를 조회합니다. "
                "KG 노드/엣지 수, 인덱싱된 문서 수와 목록, 원본 문서 수, RAG 저장소 크기를 반환합니다. "
                "사용 시점: '시스템 상태 알려줘', 'KG에 뭐가 있어?', '인덱싱된 문서가 뭐야?', "
                "'노드 몇 개야?' 등 시스템 메타정보를 물어볼 때 사용합니다. "
                "문서의 내용을 검색하려면 query_knowledge_graph를 사용하세요."
            ),
            "args": {},
        },
        {
            "name": "web_search",
            "description": (
                "DuckDuckGo를 사용하여 실시간 웹 검색을 수행합니다. "
                "인덱싱된 문서에 없는 외부 정보, 최신 뉴스, 일반 상식, 실시간 데이터가 필요할 때 사용합니다. "
                "사용 시점: '최근 뉴스 알려줘', '~에 대해 검색해줘', '~의 최신 정보', "
                "'인터넷에서 찾아봐', '~가 뭐야?' (문서에 없는 외부 지식) 등. "
                "주의: 인덱싱된 문서의 내용을 검색할 때는 query_knowledge_graph를 사용하세요. "
                "web_search는 외부 인터넷 검색 전용입니다."
            ),
            "args": {
                "query": {
                    "type": "string",
                    "description": "검색할 질문 또는 키워드. 구체적일수록 좋은 결과를 반환합니다.",
                },
                "max_results": {
                    "type": "integer",
                    "description": "반환할 최대 검색 결과 수 (기본값: 5, 최대: 10)",
                },
            },
        },
    ]


def get_utility_tools_map() -> Dict[str, Any]:
    """유틸리티 도구의 함수 매핑 반환"""
    return {
        "list_documents": list_documents,
        "get_document_info": get_document_info,
        "get_document_summary": get_document_summary,
        "get_document_content": get_document_content,
        "list_document_summaries": list_document_summaries,
        "get_system_status": get_system_status,
        "web_search": web_search,
    }
