"""
Document Converter — 다양한 문서 형식을 KG 인제스트용 텍스트 컨텍스트로 변환

지원 형식:
  .txt  / .md   → 그대로 읽기
  .pdf           → PyMuPDF (fitz) 페이지별 텍스트 추출
  .docx          → python-docx 단락 + 테이블 추출
  .xlsx          → openpyxl 시트별 → 마크다운 테이블
  .pptx          → python-pptx 슬라이드별 텍스트 + 테이블 추출

출력: 파일명 메타데이터가 포함된 구조화된 텍스트
"""

import os
from typing import Optional
from dataclasses import dataclass


@dataclass
class ConvertedDocument:
    """변환된 문서 결과."""
    filename: str
    file_type: str
    content: str
    page_count: int = 1
    error: Optional[str] = None


# ═══════════════════════════════════════════════════════
#  개별 파서
# ═══════════════════════════════════════════════════════

def _read_text(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _parse_pdf(file_path: str) -> tuple[str, int]:
    """PDF → 텍스트 (페이지 구분 포함)."""
    import fitz  # pymupdf

    doc = fitz.open(file_path)
    page_count = len(doc)
    pages = []
    for i, page in enumerate(doc):
        text = page.get_text().strip()
        if text:
            pages.append(f"--- Page {i + 1} ---\n{text}")
    doc.close()
    return "\n\n".join(pages), page_count


def _parse_docx(file_path: str) -> str:
    """DOCX → 텍스트 (단락 + 테이블)."""
    from docx import Document

    doc = Document(file_path)
    parts = []

    # 단락
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # 테이블
    for i, table in enumerate(doc.tables):
        rows_text = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_text.append(" | ".join(cells))
        if rows_text:
            parts.append(f"\n[Table {i + 1}]\n" + "\n".join(rows_text))

    return "\n\n".join(parts)


def _parse_xlsx(file_path: str) -> str:
    """XLSX → 마크다운 테이블 (시트별)."""
    import openpyxl

    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            # 빈 행 스킵
            vals = [str(c).strip() if c is not None else "" for c in row]
            if any(v for v in vals):
                rows.append(vals)

        if not rows:
            continue

        # 마크다운 테이블  
        header = " | ".join(rows[0]) if rows else ""
        sep = " | ".join(["---"] * len(rows[0])) if rows else ""
        body_rows = [" | ".join(r) for r in rows[1:]]

        table_md = f"| {header} |\n| {sep} |"
        for br in body_rows:
            table_md += f"\n| {br} |"

        parts.append(f"## Sheet: {sheet_name}\n\n{table_md}")

    wb.close()
    return "\n\n".join(parts)


def _parse_pptx(file_path: str) -> tuple[str, int]:
    """PPTX → 텍스트 (슬라이드별). python-pptx 호환 파싱."""
    from pptx import Presentation
    from lxml import etree

    prs = Presentation(file_path)
    nsmap = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}

    # 슬라이드 파트 추출 (slides iterator가 깨진 파일 대응)
    slide_parts = []
    for rel in prs.part.rels.values():
        reltype = str(rel.reltype).lower()
        if "slide" in reltype and "slideLayout" not in str(rel.reltype) and "slideMaster" not in str(rel.reltype):
            slide_parts.append(rel.target_part)

    parts = []
    for i, sp in enumerate(slide_parts):
        texts = []
        for t_elem in sp._element.findall(".//a:t", nsmap):
            if t_elem.text and t_elem.text.strip():
                texts.append(t_elem.text.strip())
        if texts:
            parts.append(f"--- Slide {i + 1} ---\n" + "\n".join(texts))

    return "\n\n".join(parts), len(slide_parts)


# ═══════════════════════════════════════════════════════
#  메인 변환 함수
# ═══════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".xlsx", ".pptx"}


def convert_document(file_path: str) -> ConvertedDocument:
    """
    단일 문서를 KG 인제스트용 텍스트 컨텍스트로 변환.
    
    출력에는 파일명 메타데이터 헤더가 포함됨:
      [Document: 파일명.확장자]
      [Type: PDF | Pages: 12]
      
      (본문 내용)
    """
    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()
    page_count = 1

    if ext not in SUPPORTED_EXTENSIONS:
        return ConvertedDocument(
            filename=filename,
            file_type=ext,
            content="",
            error=f"Unsupported file type: {ext}"
        )

    try:
        if ext in (".txt", ".md"):
            raw_content = _read_text(file_path)
            file_type = "TEXT" if ext == ".txt" else "MARKDOWN"

        elif ext == ".pdf":
            raw_content, page_count = _parse_pdf(file_path)
            file_type = "PDF"

        elif ext == ".docx":
            raw_content = _parse_docx(file_path)
            file_type = "DOCX"

        elif ext == ".xlsx":
            raw_content = _parse_xlsx(file_path)
            file_type = "XLSX"

        elif ext == ".pptx":
            raw_content, page_count = _parse_pptx(file_path)
            file_type = "PPTX"

        else:
            return ConvertedDocument(
                filename=filename, file_type=ext, content="",
                error=f"Parser not implemented for {ext}"
            )

        if not raw_content.strip():
            return ConvertedDocument(
                filename=filename, file_type=file_type, content="",
                page_count=page_count,
                error="No text content extracted"
            )

        # 메타데이터 헤더 + 본문 결합
        header = f"[Document: {filename}]\n[Type: {file_type}"
        if page_count > 1:
            header += f" | Pages: {page_count}"
        header += "]\n"

        full_content = f"{header}\n{raw_content}"

        return ConvertedDocument(
            filename=filename,
            file_type=file_type,
            content=full_content,
            page_count=page_count
        )

    except Exception as e:
        return ConvertedDocument(
            filename=filename, file_type=ext, content="",
            error=f"Parse error: {str(e)}"
        )


def convert_all_documents(docs_dir: str) -> list[ConvertedDocument]:
    """디렉토리 내 모든 지원 문서를 변환."""
    if not os.path.exists(docs_dir):
        return []

    results = []
    for entry in sorted(os.listdir(docs_dir)):
        file_path = os.path.join(docs_dir, entry)
        if not os.path.isfile(file_path):
            continue
        ext = os.path.splitext(entry)[1].lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        results.append(convert_document(file_path))

    return results


def save_contexts(docs: list[ConvertedDocument], output_dir: str) -> list[str]:
    """
    변환된 문서들을 .context.txt 파일로 저장.
    
    Returns: 저장된 파일 경로 리스트
    """
    os.makedirs(output_dir, exist_ok=True)
    saved = []

    for doc in docs:
        if doc.error or not doc.content:
            print(f"[SKIP] {doc.filename}: {doc.error}")
            continue

        base_name = os.path.splitext(doc.filename)[0]
        out_path = os.path.join(output_dir, f"{base_name}.context.txt")

        with open(out_path, "w", encoding="utf-8") as f:
            f.write(doc.content)

        saved.append(out_path)
        print(f"[SAVED] {doc.filename} → {out_path} ({len(doc.content)} chars)")

    return saved
